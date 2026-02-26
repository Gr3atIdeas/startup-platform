"""Admin bot — управление news collector через личные сообщения боту."""

import os
import logging
from datetime import datetime

from telethon import TelegramClient, events
from telethon.sessions import StringSession

import config
from storage import PostStorage

logger = logging.getLogger(__name__)

HELP_TEXT = """<b>📋 Команды админ-панели</b>

<b>Авторизация:</b>
/auth — подключить Telegram-аккаунт для парсинга
/status — текущее состояние коллектора

<b>Каналы:</b>
/channels — список отслеживаемых каналов
/add_channel &lt;id или username&gt; — добавить канал
/remove_channel &lt;id или username&gt; — удалить канал

<b>Ключевые слова (фильтр):</b>
/keywords — список ключевых слов
/add_keyword &lt;слово&gt; — добавить ключевое слово
/remove_keyword &lt;слово&gt; — удалить ключевое слово

<b>Спам-фильтр:</b>
/spam — список спам-слов
/add_spam &lt;слово&gt; — добавить спам-слово
/remove_spam &lt;слово&gt; — удалить спам-слово

<b>AI-промпт:</b>
/prompt — получить текущий промпт (.docx)
<i>Отправьте .docx файл, чтобы обновить промпт.</i>

<i>Если ключевых слов нет — пересылаются все посты.
Посты со спам-словами всегда игнорируются.</i>"""


def _parse_channel_arg(arg: str):
    """Parse channel argument — int if numeric, else string."""
    arg = arg.strip().lstrip("@")
    try:
        return int(arg)
    except ValueError:
        return arg


def create_bot_client() -> TelegramClient:
    """Create a Telethon client for the bot."""
    session_path = config.SESSION_PATH + "_bot"
    return TelegramClient(
        session_path,
        int(config.TELEGRAM_API_ID),
        config.TELEGRAM_API_HASH,
        proxy=config.get_telethon_proxy(),
    )


def register_admin_handlers(bot: TelegramClient, storage: PostStorage,
                            on_auth_complete=None, is_monitoring=None):
    """Register admin command handlers on the bot client."""

    # Auth flow state (for /auth → /code → /password sequence)
    _auth = {}

    @bot.on(events.NewMessage(pattern="/start"))
    async def on_start(event):
        if event.sender_id != config.ADMIN_ID:
            await event.reply("⛔ Доступ запрещён.")
            return
        await event.reply(HELP_TEXT, parse_mode="html")

    @bot.on(events.NewMessage(pattern="/help"))
    async def on_help(event):
        if event.sender_id != config.ADMIN_ID:
            return
        await event.reply(HELP_TEXT, parse_mode="html")

    # ── Auth ───────────────────────────────────────────────────────

    @bot.on(events.NewMessage(pattern="/auth"))
    async def on_auth(event):
        if event.sender_id != config.ADMIN_ID:
            return

        if not config.TELEGRAM_PHONE:
            await event.reply("❌ <code>TELEGRAM_PHONE</code> не задан в переменных окружения.", parse_mode="html")
            return

        # If already monitoring, no need
        if is_monitoring and is_monitoring():
            await event.reply("✅ Парсинг каналов уже активен!")
            return

        # 1) Check StringSession from env var (survives deploys)
        if config.TELETHON_SESSION:
            await event.reply("🔄 Сессия найдена (из env), пробую подключиться...")
            if on_auth_complete:
                await on_auth_complete()
            if is_monitoring and is_monitoring():
                await event.reply("✅ Парсинг каналов запущен!")
            else:
                await event.reply(
                    "⚠️ Сессия из env не работает. Удалите <code>TELETHON_SESSION</code> "
                    "в Coolify и отправьте /auth для повторной авторизации.",
                    parse_mode="html",
                )
            return

        # 2) Check file-based session (legacy, won't survive deploy)
        session_file = config.SESSION_PATH + ".session"
        if os.path.exists(session_file):
            await event.reply("🔄 Файл сессии найден, пробую подключиться...")
            if on_auth_complete:
                await on_auth_complete()
            if is_monitoring and is_monitoring():
                await event.reply("✅ Парсинг каналов запущен!")
            else:
                os.remove(session_file)
                await event.reply("⚠️ Сессия устарела, удалена. Отправьте /auth ещё раз.")
            return

        # 3) No session — start new auth flow with StringSession
        try:
            from collector import create_client
            client = create_client()  # file-based for initial auth
            await client.connect()

            result = await client.send_code_request(config.TELEGRAM_PHONE)
            _auth.update({
                "client": client,
                "phone": config.TELEGRAM_PHONE,
                "phone_code_hash": result.phone_code_hash,
            })

            phone_masked = config.TELEGRAM_PHONE[:4] + "****" + config.TELEGRAM_PHONE[-2:]
            await event.reply(
                f"📱 Код отправлен на <code>{phone_masked}</code>\n\n"
                "Введите код командой:\n"
                "<code>/code 12345</code>",
                parse_mode="html",
            )
        except Exception as e:
            logger.error("Auth start failed: %s", e)
            await event.reply(f"❌ Ошибка: {e}")
            _auth.clear()

    @bot.on(events.NewMessage(pattern=r"/code\s+(\S+)"))
    async def on_code(event):
        if event.sender_id != config.ADMIN_ID:
            return

        if "client" not in _auth:
            await event.reply("❌ Сначала отправьте /auth")
            return

        code = event.pattern_match.group(1).replace("-", "").replace(" ", "")
        client = _auth["client"]

        try:
            await client.sign_in(
                _auth["phone"],
                code,
                phone_code_hash=_auth["phone_code_hash"],
            )
        except Exception as e:
            err_name = type(e).__name__
            if "SessionPasswordNeeded" in err_name:
                _auth["need_password"] = True
                await event.reply(
                    "🔒 Аккаунт защищён 2FA.\n\n"
                    "Введите пароль:\n"
                    "<code>/password ваш_пароль</code>",
                    parse_mode="html",
                )
                return
            logger.error("Auth sign_in failed: %s", e)
            await event.reply(f"❌ Ошибка: {e}")
            await client.disconnect()
            _auth.clear()
            return

        await _finish_auth(event, client)

    @bot.on(events.NewMessage(pattern=r"/password\s+(.+)"))
    async def on_password(event):
        if event.sender_id != config.ADMIN_ID:
            return

        if not _auth.get("need_password"):
            return

        password = event.pattern_match.group(1)
        client = _auth["client"]

        try:
            await client.sign_in(password=password)
        except Exception as e:
            logger.error("2FA sign_in failed: %s", e)
            await event.reply(f"❌ Ошибка: {e}")
            await client.disconnect()
            _auth.clear()
            return

        await _finish_auth(event, client)

    async def _finish_auth(event, client):
        """Complete auth: save session, extract StringSession, start monitoring."""
        try:
            me = await client.get_me()
            name = me.first_name or "User"
            await event.reply(f"✅ Авторизован как <b>{name}</b> (id={me.id})", parse_mode="html")
        except Exception:
            pass

        # Extract StringSession for persistent storage in env var
        try:
            session_string = StringSession.save(client.session)
            await event.reply(
                "🔑 <b>ВАЖНО — сохраните сессию!</b>\n\n"
                "Скопируйте строку ниже и добавьте в Coolify как переменную:\n"
                "<code>TELETHON_SESSION</code>\n\n"
                f"<code>{session_string}</code>\n\n"
                "⚠️ Без этого при следующем деплое придётся авторизоваться заново.",
                parse_mode="html",
            )
        except Exception as e:
            logger.error("Failed to extract StringSession: %s", e)

        await client.disconnect()
        _auth.clear()

        if on_auth_complete:
            await on_auth_complete()
            if is_monitoring and is_monitoring():
                await event.reply("🚀 Парсинг каналов запущен!")
            else:
                await event.reply("⚠️ Сессия создана, но парсинг не запустился. Проверьте логи.")

    # ── Status ─────────────────────────────────────────────────────

    @bot.on(events.NewMessage(pattern="/status"))
    async def on_status(event):
        if event.sender_id != config.ADMIN_ID:
            return

        env_channels = config.SOURCE_CHANNELS_LIST
        db_channels = storage.get_channels()
        all_channels = list(set(str(c) for c in env_channels + db_channels))

        env_keywords = config.KEYWORDS_LIST
        db_keywords = storage.get_keywords()
        all_keywords = list(set(env_keywords + db_keywords))

        spam_words = storage.get_spam_words()
        processed = storage.get_processed_count()

        monitoring = is_monitoring() if is_monitoring else False
        monitoring_str = "✅ активен" if monitoring else "❌ не активен (/auth)"

        text = f"""<b>📊 Статус коллектора</b>

<b>Парсинг:</b> {monitoring_str}
<b>Каналы:</b> {len(all_channels)} ({len(env_channels)} env + {len(db_channels)} динамических)
<b>Ключевые слова:</b> {len(all_keywords)} ({len(env_keywords)} env + {len(db_keywords)} динамических)
<b>Спам-слов:</b> {len(spam_words)}
<b>Обработано постов:</b> {processed}
<b>Target:</b> {config.TARGET_GROUP_ID}
<b>Время:</b> {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"""
        await event.reply(text, parse_mode="html")

    # ── Channels ──────────────────────────────────────────────────

    @bot.on(events.NewMessage(pattern="/channels"))
    async def on_channels(event):
        if event.sender_id != config.ADMIN_ID:
            return

        env_ch = config.SOURCE_CHANNELS_LIST
        db_ch = storage.get_channels()

        lines = ["<b>📡 Каналы</b>\n"]
        if env_ch:
            lines.append("<b>Из env:</b>")
            for ch in env_ch:
                lines.append(f"  • {ch}")
        if db_ch:
            lines.append("\n<b>Динамические:</b>")
            for ch in db_ch:
                lines.append(f"  • {ch}")
        if not env_ch and not db_ch:
            lines.append("Нет каналов.")

        await event.reply("\n".join(lines), parse_mode="html")

    @bot.on(events.NewMessage(pattern=r"/add_channel\s+(.+)"))
    async def on_add_channel(event):
        if event.sender_id != config.ADMIN_ID:
            return
        arg = event.pattern_match.group(1)
        channel = _parse_channel_arg(arg)
        storage.add_channel(channel)
        await event.reply(f"✅ Канал <code>{channel}</code> добавлен.", parse_mode="html")
        logger.info("Admin added channel: %s", channel)

    @bot.on(events.NewMessage(pattern=r"/remove_channel\s+(.+)"))
    async def on_remove_channel(event):
        if event.sender_id != config.ADMIN_ID:
            return
        arg = event.pattern_match.group(1)
        channel = _parse_channel_arg(arg)
        if storage.remove_channel(channel):
            await event.reply(f"🗑 Канал <code>{channel}</code> удалён.", parse_mode="html")
            logger.info("Admin removed channel: %s", channel)
        else:
            await event.reply(f"❌ Канал <code>{channel}</code> не найден в динамических.", parse_mode="html")

    # ── Keywords ──────────────────────────────────────────────────

    @bot.on(events.NewMessage(pattern="/keywords$"))
    async def on_keywords(event):
        if event.sender_id != config.ADMIN_ID:
            return

        env_kw = config.KEYWORDS_LIST
        db_kw = storage.get_keywords()

        lines = ["<b>🔑 Ключевые слова</b>\n"]
        if env_kw:
            lines.append("<b>Из env:</b>")
            lines.append(", ".join(env_kw))
        if db_kw:
            lines.append("\n<b>Динамические:</b>")
            lines.append(", ".join(db_kw))
        if not env_kw and not db_kw:
            lines.append("Нет ключевых слов — пересылаются все посты.")

        await event.reply("\n".join(lines), parse_mode="html")

    @bot.on(events.NewMessage(pattern=r"/add_keyword\s+(.+)"))
    async def on_add_keyword(event):
        if event.sender_id != config.ADMIN_ID:
            return
        word = event.pattern_match.group(1).strip()
        storage.add_keyword(word)
        await event.reply(f"✅ Ключевое слово <code>{word.lower()}</code> добавлено.", parse_mode="html")
        logger.info("Admin added keyword: %s", word)

    @bot.on(events.NewMessage(pattern=r"/remove_keyword\s+(.+)"))
    async def on_remove_keyword(event):
        if event.sender_id != config.ADMIN_ID:
            return
        word = event.pattern_match.group(1).strip()
        if storage.remove_keyword(word):
            await event.reply(f"🗑 Ключевое слово <code>{word.lower()}</code> удалено.", parse_mode="html")
            logger.info("Admin removed keyword: %s", word)
        else:
            await event.reply(f"❌ Слово <code>{word.lower()}</code> не найдено.", parse_mode="html")

    # ── Spam words ────────────────────────────────────────────────

    @bot.on(events.NewMessage(pattern="/spam$"))
    async def on_spam(event):
        if event.sender_id != config.ADMIN_ID:
            return

        words = storage.get_spam_words()
        if words:
            text = "<b>🚫 Спам-слова</b>\n\n" + ", ".join(words)
        else:
            text = "<b>🚫 Спам-слова</b>\n\nСписок пуст."
        await event.reply(text, parse_mode="html")

    @bot.on(events.NewMessage(pattern=r"/add_spam\s+(.+)"))
    async def on_add_spam(event):
        if event.sender_id != config.ADMIN_ID:
            return
        word = event.pattern_match.group(1).strip()
        storage.add_spam_word(word)
        await event.reply(f"✅ Спам-слово <code>{word.lower()}</code> добавлено.", parse_mode="html")
        logger.info("Admin added spam word: %s", word)

    @bot.on(events.NewMessage(pattern=r"/remove_spam\s+(.+)"))
    async def on_remove_spam(event):
        if event.sender_id != config.ADMIN_ID:
            return
        word = event.pattern_match.group(1).strip()
        if storage.remove_spam_word(word):
            await event.reply(f"🗑 Спам-слово <code>{word.lower()}</code> удалено.", parse_mode="html")
            logger.info("Admin removed spam word: %s", word)
        else:
            await event.reply(f"❌ Слово <code>{word.lower()}</code> не найдено.", parse_mode="html")

    # ── Grok prompt ─────────────────────────────────────────────

    @bot.on(events.NewMessage(pattern="/prompt"))
    async def on_prompt(event):
        if event.sender_id != config.ADMIN_ID:
            return

        from grok import get_effective_prompt, SYSTEM_PROMPT
        current_prompt = get_effective_prompt(storage)
        is_custom = (current_prompt != SYSTEM_PROMPT)

        import io
        try:
            from docx import Document
        except ImportError:
            await event.reply("python-docx не установлен.")
            return

        doc = Document()
        doc.add_paragraph(current_prompt)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = "grok_prompt.docx"

        status = "из БД" if is_custom else "по умолчанию"
        await bot.send_file(
            event.chat_id,
            buf,
            caption=f"Текущий промпт ({status}, {len(current_prompt)} симв.).\n"
                    "Отправьте .docx файл, чтобы обновить.",
        )

    @bot.on(events.NewMessage())
    async def on_docx_upload(event):
        if event.sender_id != config.ADMIN_ID:
            return
        if not event.message.document:
            return

        file_name = ""
        for attr in event.message.document.attributes:
            if hasattr(attr, "file_name") and attr.file_name:
                file_name = attr.file_name
                break

        if not file_name.lower().endswith(".docx"):
            return

        data = await event.message.download_media(bytes)
        if not data:
            await event.reply("Не удалось скачать файл.")
            return

        import io
        try:
            from docx import Document
        except ImportError:
            await event.reply("python-docx не установлен.")
            return

        try:
            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            await event.reply(f"Ошибка чтения .docx: {e}")
            return

        if not text.strip():
            await event.reply("Документ пустой. Промпт не изменён.")
            return

        storage.set_setting("grok_prompt", text.strip())

        from html import escape
        preview = escape(text.strip()[:200])
        if len(text.strip()) > 200:
            preview += "..."

        await event.reply(
            f"Промпт обновлён ({len(text.strip())} симв.).\n\n"
            f"<b>Превью:</b>\n<code>{preview}</code>",
            parse_mode="html",
        )
        logger.info("Admin updated Grok prompt (%d chars)", len(text.strip()))

    logger.info("Admin bot handlers registered (admin_id=%d)", config.ADMIN_ID)


async def setup_bot_commands(bot: TelegramClient):
    """Set bot commands visible in Telegram's '/' menu."""
    try:
        from telethon.tl.functions.bots import SetBotCommandsRequest
        from telethon.tl.types import BotCommand, BotCommandScopeDefault
    except ImportError:
        logger.warning("Could not import BotCommand types, skipping menu setup")
        return

    commands = [
        BotCommand(command="help", description="Все команды"),
        BotCommand(command="status", description="Статус коллектора"),
        BotCommand(command="channels", description="Список каналов"),
        BotCommand(command="add_channel", description="Добавить канал"),
        BotCommand(command="remove_channel", description="Удалить канал"),
        BotCommand(command="keywords", description="Ключевые слова"),
        BotCommand(command="add_keyword", description="Добавить ключевое слово"),
        BotCommand(command="remove_keyword", description="Удалить ключевое слово"),
        BotCommand(command="spam", description="Спам-слова"),
        BotCommand(command="add_spam", description="Добавить спам-слово"),
        BotCommand(command="remove_spam", description="Удалить спам-слово"),
        BotCommand(command="prompt", description="Промпт Grok (.docx)"),
        BotCommand(command="auth", description="Авторизация Telegram"),
    ]

    try:
        await bot(SetBotCommandsRequest(
            scope=BotCommandScopeDefault(),
            lang_code="",
            commands=commands,
        ))
        logger.info("Bot commands menu set (%d commands)", len(commands))
    except Exception as e:
        logger.error("Failed to set bot commands: %s", e)
