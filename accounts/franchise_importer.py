"""
AI-powered franchise importer.

Parses uploaded files (PDF, DOCX, Excel, images) and URLs,
extracts text, sends to Grok AI for structured data extraction,
and creates a franchise draft with pre-filled fields.
"""

import io
import json
import logging
import re

import requests as http_requests
from django.conf import settings

logger = logging.getLogger(__name__)


# ── File Parsers ─────────────────────────────────────────────


def parse_pdf(file_obj):
    """Extract text from PDF file."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages[:30]:  # limit to 30 pages
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                # Also extract tables
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            text_parts.append(" | ".join(str(cell or "") for cell in row))
        return "\n\n".join(text_parts)
    except ImportError:
        logger.warning("pdfplumber not installed, trying PyPDF2")
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_obj)
            return "\n\n".join(
                page.extract_text() or "" for page in reader.pages[:30]
            )
        except ImportError:
            logger.error("No PDF parser available (install pdfplumber or PyPDF2)")
            return ""
    except Exception as e:
        logger.error("PDF parse error: %s", e)
        return ""


def parse_docx(file_obj):
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_obj)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except ImportError:
        logger.error("python-docx not installed")
        return ""
    except Exception as e:
        logger.error("DOCX parse error: %s", e)
        return ""


def parse_excel(file_obj):
    """Extract text from Excel file."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_obj, read_only=True, data_only=True)
        parts = []
        for sheet in wb.sheetnames[:5]:
            ws = wb[sheet]
            parts.append(f"=== Лист: {sheet} ===")
            for row in ws.iter_rows(max_row=100, values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except Exception as e:
        logger.error("Excel parse error: %s", e)
        return ""


def parse_url(url):
    """Scrape text content from a URL."""
    try:
        from bs4 import BeautifulSoup
        headers = {
            "User-Agent": "Mozilla/5.0 (compatible; GreatIdeasBot/1.0)"
        }
        proxy_url = getattr(settings, 'PROXY_URL', '')
        proxies = {'http': proxy_url, 'https': proxy_url} if proxy_url else None

        resp = http_requests.get(url, headers=headers, timeout=15, proxies=proxies)
        resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "html.parser")

        # Remove scripts, styles, nav
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        # Collapse multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text[:15000]  # limit
    except ImportError:
        logger.error("beautifulsoup4 not installed")
        return ""
    except Exception as e:
        logger.error("URL scrape error for %s: %s", url, e)
        return ""


def parse_file(file_obj, filename):
    """Route file to appropriate parser based on extension."""
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

    if ext == 'pdf':
        return parse_pdf(file_obj)
    elif ext in ('docx', 'doc'):
        return parse_docx(file_obj)
    elif ext in ('xlsx', 'xls'):
        return parse_excel(file_obj)
    elif ext in ('txt', 'csv'):
        try:
            content = file_obj.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            return content[:15000]
        except Exception as e:
            logger.error("Text file parse error: %s", e)
            return ""
    else:
        return ""  # Images and other files — no text to extract


def is_image_file(filename):
    """Check if file is an image."""
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    return ext in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp', 'svg')


# ── AI Extraction ────────────────────────────────────────────


EXTRACTION_PROMPT = """Ты — система извлечения данных о франшизах для платформы GreatIdeas.ru.

На основе предоставленного текста заполни JSON-структуру франшизы.
Извлеки ВСЕ данные, которые можешь найти. Если данных нет — оставь null.
НЕ ВЫДУМЫВАЙ данные. Если информации нет в тексте — ставь null.

Верни СТРОГО JSON (без markdown-обёртки):
{
  "title": "Название франшизы (string)",
  "short_description": "Краткое описание, 1-2 предложения (string)",
  "description": "Полное описание франшизы, HTML-форматирование с <p>, <h3>, <ul> (string)",
  "terms": "Условия сотрудничества, HTML (string или null)",
  "investment_size": число в рублях или null,
  "franchise_cost": паушальный взнос в рублях или null,
  "payback_period": срок окупаемости в месяцах (число) или null,
  "profit_calculation": "Расчёт прибыли, HTML (string или null)",
  "direction": "Категория: одна из (Technology, Healthcare, Finance, Beauty, Cafe, Delivery, Fastfood, Sport, Education, Entertainment) или null",
  "contact_website": "URL сайта или null",
  "contact_telegram": "Telegram контакт или null",
  "contact_whatsapp": "WhatsApp/телефон или null",
  "own_businesses_count": число собственных точек или null,
  "franchise_businesses_count": число франчайзинговых точек или null,
  "cities": ["список городов присутствия"] или [],
  "missing_fields": ["список полей, которые не удалось найти в тексте"]
}
"""


def extract_franchise_data(all_text):
    """Send combined text to Grok AI and extract structured franchise data."""
    api_key = getattr(settings, 'GROK_API_KEY', '')
    model = getattr(settings, 'GROK_MODEL', 'grok-3-mini')

    if not api_key:
        logger.error("GROK_API_KEY not configured")
        return None

    # Truncate if too long
    if len(all_text) > 25000:
        all_text = all_text[:25000] + "\n\n[...текст обрезан...]"

    proxy_url = getattr(settings, 'PROXY_URL', '')
    proxies = {'http': proxy_url, 'https': proxy_url} if proxy_url else None

    try:
        response = http_requests.post(
            "https://api.x.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": model,
                "messages": [
                    {"role": "system", "content": EXTRACTION_PROMPT},
                    {"role": "user", "content": f"Извлеки данные о франшизе из следующего текста:\n\n{all_text}"},
                ],
                "temperature": 0.3,  # Very low — we want factual extraction
                "max_tokens": 4000,
            },
            proxies=proxies or None,
            timeout=90,
        )
        response.raise_for_status()
        data = response.json()
        content = data["choices"][0]["message"]["content"]

        # Parse JSON
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0].strip()
        elif "```" in content:
            content = content.split("```")[1].split("```")[0].strip()

        return json.loads(content)

    except http_requests.RequestException as e:
        logger.error("Grok extraction API error: %s", e)
        return None
    except (json.JSONDecodeError, KeyError, IndexError) as e:
        logger.error("Failed to parse Grok extraction response: %s", e)
        return None


# ── Franchise Creator ────────────────────────────────────────


def create_franchise_from_data(extracted, image_files=None, user=None):
    """Create a Franchise object from extracted AI data. Returns (franchise, warnings)."""
    from .models import Franchises, Directions, City, FranchiseLocation
    from django.utils import timezone
    import uuid

    warnings = []
    data = extracted or {}

    title = data.get('title') or ''
    if not title:
        warnings.append("Название не определено — введите вручную")
        title = "Новая франшиза (импорт)"

    # Find direction
    direction = None
    direction_name = data.get('direction')
    if direction_name:
        direction = Directions.objects.filter(direction_name__iexact=direction_name).first()
        if not direction:
            direction = Directions.objects.filter(direction_name__icontains=direction_name[:4]).first()
        if not direction:
            warnings.append(f"Категория '{direction_name}' не найдена — выберите вручную")

    franchise = Franchises(
        title=title,
        short_description=data.get('short_description') or '',
        description=data.get('description') or '',
        terms=data.get('terms') or '',
        investment_size=data.get('investment_size'),
        franchise_cost=data.get('franchise_cost'),
        payback_period=data.get('payback_period'),
        profit_calculation=data.get('profit_calculation') or '',
        direction=direction,
        contact_website=data.get('contact_website') or '',
        contact_telegram=data.get('contact_telegram') or '',
        contact_whatsapp=data.get('contact_whatsapp') or '',
        own_businesses_count=data.get('own_businesses_count') or 0,
        franchise_businesses_count=data.get('franchise_businesses_count') or 0,
        status='pending',
        owner=user,
        created_at=timezone.now(),
        updated_at=timezone.now(),
    )
    franchise.save()

    # Handle images
    if image_files:
        from .utils import upload_file_to_s3_sync
        logo_ids = []
        creative_ids = []
        for i, img_file in enumerate(image_files[:11]):  # max 10 creatives + 1 logo
            file_id = str(uuid.uuid4())
            try:
                img_file.seek(0)
                file_data = img_file.read()
                content_type = getattr(img_file, 'content_type', 'image/jpeg')
                file_type = 'logo' if i == 0 else 'creative'
                upload_file_to_s3_sync(
                    file_data=file_data,
                    file_name=img_file.name,
                    content_type=content_type,
                    entity_type_name='franchise',
                    entity_id=franchise.franchise_id,
                    file_type_name=file_type,
                    original_filename=img_file.name,
                    file_id=file_id,
                )
                if i == 0:
                    logo_ids.append(file_id)
                else:
                    creative_ids.append(file_id)
            except Exception as e:
                warnings.append(f"Не удалось загрузить {img_file.name}: {e}")

        franchise.logo_urls = logo_ids
        franchise.creatives_urls = creative_ids
        if creative_ids:
            franchise.slider_images = creative_ids[:4]
        franchise.save()

    # Handle cities
    cities_list = data.get('cities') or []
    created_locations = 0
    for city_name in cities_list:
        city_name = city_name.strip()
        if not city_name:
            continue
        city = City.objects.filter(name__iexact=city_name).first()
        if not city:
            # Try partial match
            city = City.objects.filter(name__icontains=city_name[:5]).first()
        if city:
            FranchiseLocation.objects.get_or_create(
                franchise=franchise, city=city,
                defaults={'status': 'active'}
            )
            created_locations += 1
        else:
            warnings.append(f"Город '{city_name}' не найден в справочнике")

    if created_locations:
        logger.info("Created %d locations for franchise %d", created_locations, franchise.franchise_id)

    # Report missing fields
    missing = data.get('missing_fields') or []
    if missing:
        warnings.extend([f"Не найдено: {field}" for field in missing])

    # Check what's still empty
    if not franchise.short_description:
        warnings.append("Короткое описание не заполнено")
    if not franchise.investment_size:
        warnings.append("Размер инвестиций не определён")
    if not franchise.franchise_cost:
        warnings.append("Паушальный взнос не определён")
    if not franchise.logo_urls:
        warnings.append("Логотип не загружен")

    return franchise, warnings
